from docx import Document
import os


def merge_docs(output_file):
    current_directory = os.getcwd()

    merged_document = Document()
    for fileName in os.listdir(current_directory):
        if fileName.endswith(".docx") and fileName != output_file:
            file_path = os.path.join(current_directory, fileName)
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                merged_document.add_paragraph(paragraph.text)

            merged_document.add_page_break()

    merged_document.save(output_file)
    print(f"Document {output_file} created successfully!")


output_file = 'merged_output.docx'

# Call the merge function
merge_docs(output_file)
